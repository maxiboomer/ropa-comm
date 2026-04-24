#!/usr/bin/env python3
"""
Gerador do Roadmap de Implantação — RoPA-WEB Instituição
Produz o arquivo DOCX institucional completo.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path

# ── Paleta de cores ───────────────────────────────────────────────────────────
AZUL        = RGBColor(0x1F, 0x3D, 0x7A)
AZUL_CLARO  = RGBColor(0xEE, 0xF2, 0xFF)
AZUL_MEDIO  = RGBColor(0xD0, 0xD9, 0xF0)
CINZA       = RGBColor(0x5A, 0x5A, 0x5A)
CINZA_CLARO = RGBColor(0xF5, 0xF5, 0xF5)
VERDE       = RGBColor(0x19, 0x87, 0x54)
VERMELHO    = RGBColor(0xDC, 0x35, 0x45)
LARANJA     = RGBColor(0xFD, 0x7E, 0x14)
AMARELO     = RGBColor(0xE9, 0xB9, 0x49)
BRANCO      = RGBColor(0xFF, 0xFF, 0xFF)
PRETO       = RGBColor(0x22, 0x22, 0x22)

AZUL_HEX        = "1F3D7A"
AZUL_CLARO_HEX  = "EEF2FF"
CINZA_CLARO_HEX = "F5F5F5"
AZUL_MEDIO_HEX  = "D0D9F0"
BRANCO_HEX      = "FFFFFF"
CINZA_HEX       = "5A5A5A"

# ── Helpers XML ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))  # cm to twips (approx)
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def add_paragraph_border_bottom(paragraph, color="1F3D7A", size=12, space=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_paragraph_border_top(paragraph, color="D0D9F0", size=4, space=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), str(space))
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


# ── Document helpers ─────────────────────────────────────────────────────────

def styled_run(paragraph, text, bold=False, italic=False, size=11, color=None, font="Arial"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1, add_border=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    size = 16 if level == 1 else (13 if level == 2 else 11)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = AZUL
    if level == 1 and add_border:
        add_paragraph_border_bottom(p, color=AZUL_HEX, size=8)
    return p


def add_para(doc, text, bold=False, italic=False, size=10, color=None, indent=False, space_before=4, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = color or PRETO
    return p


def add_bullet(doc, text, size=10, marker="•", indent_cm=0.8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{marker}  {text}")
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = PRETO
    return p


def add_check(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run_box = p.add_run("☐  ")
    run_box.font.size = Pt(size)
    run_box.font.name = "Arial"
    run_box.font.color.rgb = AZUL
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = PRETO
    return p


def add_space(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(size)
    p.paragraph_format.space_after = Pt(0)


def make_header_cell(cell, text, size=9.5, bg=AZUL_HEX, text_color=BRANCO, bold=True):
    set_cell_bg(cell, bg)
    set_cell_borders(cell, color=AZUL_HEX, size="4")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(text_color) if isinstance(text_color, str) else text_color
    return cell


def make_data_cell(cell, text, bg=None, text_color=None, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    if bg:
        set_cell_bg(cell, bg)
    set_cell_borders(cell)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = text_color or PRETO
    return cell


def add_code_block(doc, lines):
    """Bloco de código estilo terminal."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1E1E1E")
    set_cell_borders(cell, color="444444")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        color = RGBColor(0x6A, 0x99, 0x55) if line.startswith("#") else RGBColor(0xFF, 0xFF, 0xFF)
        run = p.add_run(line)
        run.font.size = Pt(8.5)
        run.font.name = "Courier New"
        run.font.color.rgb = color
    return table


def add_phase_block(doc, num, title, duration, color_hex, items):
    """Bloco visual de fase com cor lateral."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # Coluna colorida (fase)
    set_cell_bg(left_cell, color_hex)
    set_cell_borders(left_cell, color=color_hex)
    set_cell_margins(left_cell, top=160, bottom=160, left=160, right=160)
    left_cell._tc.get_or_add_tcPr()
    w_tcW = OxmlElement("w:tcW")
    w_tcW.set(qn("w:w"), "2200")
    w_tcW.set(qn("w:type"), "dxa")
    left_cell._tc.tcPr.append(w_tcW)

    p1 = left_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"FASE {num}")
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.name = "Arial"
    r1.font.color.rgb = BRANCO

    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    r2 = p2.add_run(duration)
    r2.font.size = Pt(9)
    r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor(0xDD, 0xDD, 0xFF)

    # Coluna de conteúdo
    right_cell._tc.get_or_add_tcPr()
    w_tcW2 = OxmlElement("w:tcW")
    w_tcW2.set(qn("w:w"), "7500")
    w_tcW2.set(qn("w:type"), "dxa")
    right_cell._tc.tcPr.append(w_tcW2)

    set_cell_bg(right_cell, AZUL_HEX)
    set_cell_borders(right_cell, color=AZUL_HEX)
    set_cell_margins(right_cell, top=120, bottom=120, left=200, right=160)

    p_title = right_cell.paragraphs[0]
    r_title = p_title.add_run(title)
    r_title.bold = True
    r_title.font.size = Pt(13)
    r_title.font.name = "Arial"
    r_title.font.color.rgb = BRANCO

    # Items subcoluna
    for item in items:
        pi = right_cell.add_paragraph()
        pi.paragraph_format.space_before = Pt(3)
        pi.paragraph_format.space_after = Pt(2)
        r_arrow = pi.add_run("▸  ")
        r_arrow.font.size = Pt(9)
        r_arrow.font.name = "Arial"
        r_arrow.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)
        r_item = pi.add_run(item)
        r_item.font.size = Pt(9)
        r_item.font.name = "Arial"
        r_item.font.color.rgb = BRANCO


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Margens ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Estilos default ───────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════════════════════════════════════════

add_space(doc, 48)

p_inst = doc.add_paragraph()
p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_inst.paragraph_format.space_after = Pt(4)
add_paragraph_border_bottom(p_inst, color=AZUL_HEX, size=10)
r_inst = p_inst.add_run("INSTITUIÇÃO")
r_inst.bold = True
r_inst.font.size = Pt(14)
r_inst.font.name = "Arial"
r_inst.font.color.rgb = AZUL

p_algpd = doc.add_paragraph()
p_algpd.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_algpd.paragraph_format.space_after = Pt(48)
r_algpd = p_algpd.add_run("Unidade de Proteção de Dados")
r_algpd.font.size = Pt(11)
r_algpd.font.name = "Arial"
r_algpd.font.color.rgb = CINZA

add_space(doc, 24)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(6)
r_t = p_title.add_run("ROADMAP DE IMPLANTAÇÃO")
r_t.bold = True
r_t.font.size = Pt(28)
r_t.font.name = "Arial"
r_t.font.color.rgb = AZUL

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(6)
r_s = p_sub.add_run("Sistema RoPA-WEB")
r_s.bold = True
r_s.font.size = Pt(20)
r_s.font.name = "Arial"
r_s.font.color.rgb = CINZA

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_desc.paragraph_format.space_after = Pt(48)
r_d = p_desc.add_run("Registro de Atividades de Tratamento na Infraestrutura Instituição")
r_d.italic = True
r_d.font.size = Pt(11)
r_d.font.name = "Arial"
r_d.font.color.rgb = CINZA

add_space(doc, 24)

# Tabela de metadados da capa
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Versão", "1.0 — Abril/2026"),
    ("Encarregado(a) (DPO)", "Encarregado(a) de Proteção de Dados"),
    ("Unidade Responsável", "Unidade de Proteção de Dados"),
    ("Classificação", "Uso Interno — Instituição"),
]
for i, (label, value) in enumerate(meta_data):
    bg = AZUL_CLARO_HEX if i % 2 == 0 else CINZA_CLARO_HEX
    make_header_cell(meta_table.cell(i, 0), label, bg=AZUL_HEX, size=9)
    make_data_cell(meta_table.cell(i, 1), value, bg=bg, size=9)
    set_row_height(meta_table.rows[i], 0.8)

add_space(doc, 48)

p_norm = doc.add_paragraph()
p_norm.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_n = p_norm.add_run("Conforme LGPD — Lei 13.709/2018, Art. 37  |  Resolução TSE 23.222/2010")
r_n.italic = True
r_n.font.size = Pt(9)
r_n.font.name = "Arial"
r_n.font.color.rgb = CINZA

# Quebra de página
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. APRESENTAÇÃO
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. Apresentação", level=1)

add_para(doc, (
    "Este documento define o roadmap completo para a implantação do sistema RoPA-WEB na "
    "infraestrutura do Instituição, compreendendo todas as etapas "
    "técnicas, administrativas e de segurança necessárias para colocar a solução em operação em ambiente "
    "de produção."
))
add_para(doc, (
    "O RoPA-WEB é uma aplicação web desenvolvida em Python/Flask para gestão do Registro de Atividades "
    "de Tratamento (RoPA) do Instituição, atendendo às exigências do Art. 37 da Lei Geral de Proteção de "
    "Dados (LGPD — Lei 13.709/2018). A solução substitui processos manuais em planilhas e documentos "
    "dispersos por um sistema centralizado, auditável e exportável em formatos institucionais (XLSX, PDF, JSON)."
))
add_para(doc, "Este documento é dirigido às seguintes unidades:", bold=True)
add_bullet(doc, "Unidade de TI — Seção de Tecnologia da Informação (implantação técnica e infraestrutura)")
add_bullet(doc, "Unidade de Proteção de Dados (gestão e operação do sistema)")
add_bullet(doc, "SEGINF — Segurança da Informação (análise de riscos e conformidade)")
add_bullet(doc, "SECAD — Assessoria de Administração (aprovação e recursos)")

add_space(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 2. BASE NORMATIVA
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. Base Normativa", level=1)
add_para(doc, "A implantação e operação do RoPA-WEB observa o seguinte conjunto normativo:")
add_space(doc)

normas = [
    ("LGPD — Lei 13.709/2018", "Arts. 37 (registro obrigatório), 46 (segurança), 5º II (dados sensíveis), 7º e 11 (bases legais)"),
    ("Resolução TSE 23.222/2010", "Tabela de temporalidade e destinação de documentos da Justiça Eleitoral"),
    ("Res. Instituição nº 971/2026", "Regulamento da Secretaria do Instituição e atribuições das unidades"),
    ("Portaria Instituição nº 302/2025", "Política de Segurança da Informação do Instituição"),
    ("ABNT NBR ISO/IEC 27001", "Sistemas de gestão de segurança da informação — Requisitos"),
    ("IN SGD/ME nº 1/2019", "Disciplina o processo de contratação de TIC na Administração Pública Federal"),
]
t = doc.add_table(rows=len(normas) + 1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
make_header_cell(t.cell(0, 0), "Norma")
make_header_cell(t.cell(0, 1), "Dispositivos Aplicáveis")
for i, (norma, dispositivos) in enumerate(normas, 1):
    bg = AZUL_CLARO_HEX if i % 2 == 1 else CINZA_CLARO_HEX
    make_data_cell(t.cell(i, 0), norma, bg=bg, bold=True, size=9)
    make_data_cell(t.cell(i, 1), dispositivos, bg=bg, size=9)

add_space(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. ARQUITETURA
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. Arquitetura do Sistema", level=1)

add_heading(doc, "3.1 Stack Tecnológica", level=2)
stack = [
    ("Backend / API", "Python 3.10+ / Flask 3.x", "Linguagem amplamente adotada na adm. pública; sem custo de licença"),
    ("Servidor WSGI", "Gunicorn 21+", "Servidor de produção estável e amplamente testado para Flask"),
    ("Proxy Reverso", "Nginx 1.24+", "TLS termination, rate limiting, compressão gzip, logs de acesso"),
    ("Banco de Dados", "SQLite 3.x → PostgreSQL 15+", "SQLite fase inicial (baixo volume); migração para PostgreSQL na escala"),
    ("Frontend", "HTML5 + Bootstrap 5", "Sem dependência de SPA; acessível; funciona sem JS avançado"),
    ("Exportação", "openpyxl + ReportLab", "Geração de XLSX e PDF institucionais sem licença adicional"),
    ("Container (opc.)", "Docker 24+ / Compose 2.x", "Isolamento de dependências; facilita rollback e replicação"),
]
t = doc.add_table(rows=len(stack) + 1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
make_header_cell(t.cell(0, 0), "Camada")
make_header_cell(t.cell(0, 1), "Tecnologia")
make_header_cell(t.cell(0, 2), "Justificativa")
for i, (camada, tech, just) in enumerate(stack, 1):
    bg = AZUL_CLARO_HEX if i % 2 == 1 else CINZA_CLARO_HEX
    make_data_cell(t.cell(i, 0), camada, bg=bg, bold=True, size=9)
    make_data_cell(t.cell(i, 1), tech, bg=bg, size=9)
    make_data_cell(t.cell(i, 2), just, bg=bg, size=9)

add_space(doc)
add_heading(doc, "3.2 Topologia de Rede Proposta", level=2)
add_para(doc, (
    "A solução opera integralmente na rede interna do Instituição (Intranet), sem exposição à Internet pública. "
    "A arquitetura de três camadas abaixo é a recomendada:"
))
add_space(doc, 4)

topo_lines = [
    "[ Navegador do Usuário — Rede Interna Instituição ]",
    "         ↓  HTTPS / TLS 1.3",
    "[ Nginx — Proxy Reverso + TLS Termination + Rate Limiting ]",
    "         ↓  HTTP Unix Socket (loopback)",
    "[ Gunicorn — WSGI App Server  (2–4 workers) ]",
    "         ↓  Python WSGI",
    "[ Flask App — Lógica de Negócio + Jinja2 Templates ]",
    "         ↓  SQLite3 / PostgreSQL (via psycopg2)",
    "[ Banco de Dados — Volume Persistente + Backup NAS ]",
]
topo_table = doc.add_table(rows=1, cols=1)
topo_table.alignment = WD_TABLE_ALIGNMENT.LEFT
c = topo_table.cell(0, 0)
set_cell_bg(c, AZUL_CLARO_HEX)
set_cell_borders(c, color=AZUL_HEX)
set_cell_margins(c, top=120, bottom=120, left=200, right=200)
c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
for line in topo_lines:
    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    is_arrow = "↓" in line
    r = p.add_run(line)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.bold = not is_arrow
    r.font.color.rgb = CINZA if is_arrow else AZUL

add_space(doc)
add_heading(doc, "3.3 Requisitos de Hardware", level=2)
hw = [
    ("CPU",           "2 vCPUs",              "4 vCPUs (suporte a múltiplos workers Gunicorn)"),
    ("RAM",           "2 GB",                 "4 GB (geração de PDF/XLSX simultâneos)"),
    ("Disco",         "20 GB",                "50 GB SSD (logs, banco, exports, backups locais)"),
    ("S.O.",          "Ubuntu 22.04 LTS",     "Ubuntu 22.04 LTS ou RHEL 9 (padrão Instituição)"),
    ("Rede",          "100 Mbps Intranet",    "1 Gbps; IP fixo; FQDN: ropa.intranet.example.org"),
]
t = doc.add_table(rows=len(hw) + 1, cols=3)
make_header_cell(t.cell(0, 0), "Recurso")
make_header_cell(t.cell(0, 1), "Mínimo")
make_header_cell(t.cell(0, 2), "Recomendado (produção)")
for i, (r, mi, rec) in enumerate(hw, 1):
    bg = AZUL_CLARO_HEX if i % 2 == 1 else CINZA_CLARO_HEX
    make_data_cell(t.cell(i, 0), r, bg=bg, bold=True, size=9)
    make_data_cell(t.cell(i, 1), mi, bg=bg, size=9)
    make_data_cell(t.cell(i, 2), rec, bg=bg, size=9)

add_space(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. PRÉ-REQUISITOS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. Pré-requisitos de Infraestrutura", level=1)

add_heading(doc, "4.1 Softwares e Dependências", level=2)
prereqs = [
    "Python 3.10 ou superior  (verificar: python3 --version)",
    "pip 23+ e venv nativo  (sudo apt install python3-venv)",
    "Nginx 1.24+  (sudo apt install nginx)",
    "Gunicorn 21+  (instalado via pip no ambiente virtual)",
    "SQLite 3.35+  (incluso no Ubuntu 22.04 LTS)",
    "Git 2.34+  (clone e controle de versão do código-fonte)",
    "Certbot + Let's Encrypt OU certificado da CA interna Instituição",
    "Docker 24+ e Docker Compose 2.20+  (opcional — instalação em container)",
    "pip-audit  (auditoria de vulnerabilidades nas dependências Python)",
]
for item in prereqs:
    add_bullet(doc, item)

add_space(doc)
add_heading(doc, "4.2 Contas e Permissões Necessárias", level=2)
perms = [
    ("Usuário 'ropaapp' (sem shell)", "Unidade de TI / Adm. Linux", "Executar o processo Gunicorn com privilégio mínimo"),
    ("Sudo restrito (Nginx / systemd)", "Unidade de TI", "Reiniciar serviços em atualizações e rollback"),
    ("Acesso ao GitLab Instituição", "Unidade de TI + LGPD", "Clone, deploy e controle de versão do código"),
    ("Regra de firewall (porta 443)", "SEGINF / Infra", "Acesso HTTPS pela rede interna do Instituição"),
    ("Certificado TLS (FQDN interno)", "SEGINF / CA interna", "Comunicação criptografada (HTTPS obrigatório)"),
    ("Acesso ao NAS / storage", "Unidade de TI", "Armazenamento dos backups automáticos diários"),
]
t = doc.add_table(rows=len(perms) + 1, cols=3)
make_header_cell(t.cell(0, 0), "Conta / Permissão")
make_header_cell(t.cell(0, 1), "Responsável")
make_header_cell(t.cell(0, 2), "Finalidade")
for i, (perm, resp, fin) in enumerate(perms, 1):
    bg = AZUL_CLARO_HEX if i % 2 == 1 else CINZA_CLARO_HEX
    make_data_cell(t.cell(i, 0), perm, bg=bg, bold=True, size=9)
    make_data_cell(t.cell(i, 1), resp, bg=bg, size=9)
    make_data_cell(t.cell(i, 2), fin, bg=bg, size=9)

add_space(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. FASES DE IMPLANTAÇÃO
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Fases de Implantação", level=1)
add_para(doc, (
    "O roadmap está dividido em cinco fases sequenciais com critérios de entrada e saída definidos. "
    "Nenhuma fase deve ser iniciada sem a conclusão formal da fase anterior."
))
add_space(doc)

add_phase_block(doc, "0", "Preparação e Planejamento", "Semana 1", AZUL_HEX, [
    "Reunião de kickoff com Unidade de TI, LGPD e SEGINF — definir cronograma e responsáveis",
    "Solicitar provisionamento da VM conforme especificações da seção 3.3",
    "Solicitar abertura de regra de firewall para porta 443 na rede interna",
    "Criar repositório Git no GitLab Instituição e conceder acessos às equipes",
    "Solicitar certificado TLS à CA interna para ropa.intranet.example.org",
    "Elaborar Plano de Comunicação para os usuários-chave (DPO + cadastradores)",
    "Definir política de backup: frequência diária, retenção 30 dias, destino NAS",
    "Agendar janela de manutenção para deploy em homologação e produção",
])
add_space(doc, 12)

add_phase_block(doc, "1", "Ambiente de Homologação", "Semanas 2–3", "2563EB", [
    "Instalar Ubuntu Server 22.04 LTS na VM de homologação; aplicar patches de segurança",
    "Criar usuário de sistema 'ropaapp' sem shell de login (adduser --system --no-create-home)",
    "Clonar repositório Git e instalar dependências no ambiente virtual Python (.venv)",
    "Configurar Gunicorn com 2 workers e socket Unix (/run/ropa/ropa.sock)",
    "Configurar Nginx como proxy reverso com TLS 1.3 e headers de segurança HTTP",
    "Configurar serviços systemd (ropa.service) com restart automático em falha",
    "Popular banco de homologação com dados de exemplo (python app.py → seed)",
    "Executar testes funcionais: CRUD completo, exportação XLSX/PDF, validação LGPD",
    "Revisar logs de acesso Nginx e logs da aplicação em busca de erros ou anomalias",
    "Obter aceite formal da LGPD sobre funcionalidades testadas em homologação",
])
add_space(doc, 12)

add_phase_block(doc, "2", "Segurança e Hardening", "Semana 4", "7C3AED", [
    "Auditoria Nginx: HSTS, X-Content-Type-Options, X-Frame-Options, CSP, Referrer-Policy",
    "Verificar FLASK_SECRET_KEY: string aleatória ≥ 32 bytes, nunca estática no código",
    "Configurar rate limiting no Nginx (limit_req_zone) para mitigar força bruta",
    "Revisar permissões de arquivo: banco SQLite (chmod 600) e diretório exports (700)",
    "Ativar UFW: permitir somente portas 22 (SSH restrito por IP) e 443 (HTTPS interno)",
    "Configurar logrotate para logs Nginx e Gunicorn (retenção mínima de 90 dias)",
    "Executar pip-audit para identificar CVEs nas dependências Python instaladas",
    "Elaborar RIPD (Relatório de Impacto) conforme Art. 10, §3 LGPD",
    "Aprovação formal do SEGINF (registro em processo SEI) para prosseguir à produção",
])
add_space(doc, 12)

add_phase_block(doc, "3", "Implantação em Produção", "Semanas 5–6", "059669", [
    "Provisionar VM de produção com as mesmas especificações de homologação",
    "Repetir Fases 1 e 2 no ambiente de produção seguindo o runbook documentado",
    "Configurar backup automático diário: sqlite3 .backup para NAS Instituição",
    "Migrar dados validados da homologação para produção (se aplicável)",
    "Executar smoke tests: dashboard, criação de atividade, exportação PDF institucional",
    "Comunicar URL de produção (ropa.intranet.example.org) aos usuários autorizados",
    "Realizar treinamento presencial ou por videoconferência com as unidades cadastradoras",
    "Registrar sistema RoPA-WEB no inventário de ativos de TI do Instituição",
    "Emitir Termo de Aceite de Implantação assinado por LGPD, Unidade de TI e SEGINF",
])
add_space(doc, 12)

add_phase_block(doc, "4", "Operação e Monitoramento Contínuo", "A partir da Semana 7", "D97706", [
    "Monitorar disponibilidade via ferramenta de monitoramento Instituição (ex: Zabbix/Grafana)",
    "Revisão mensal do RoPA pela LGPD — meta: 100% das atividades com completude ≥ 80%",
    "Aplicar patches de segurança do SO e dependências Python trimestralmente",
    "Executar restore test de backup semestralmente (validar integridade dos dados)",
    "Avaliar migração de SQLite para PostgreSQL: critério ≥ 500 atividades ou 10+ usuários",
    "Revisar e atualizar RIPD anualmente ou a cada mudança funcional significativa",
    "Revisar acessos e permissões de usuário semestralmente (auditoria de acessos)",
    "Manter log de auditoria de alterações por mínimo 5 anos — Res. TSE 23.222/2010",
])

add_space(doc, 16)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 6. CONFIGURAÇÕES TÉCNICAS DETALHADAS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "6. Configurações Técnicas Detalhadas", level=1)

add_heading(doc, "6.1 Serviço Systemd — ropa.service", level=2)
add_para(doc, "Criar o arquivo /etc/systemd/system/ropa.service:")
add_code_block(doc, [
    "[Unit]",
    "Description=RoPA Instituição — Gunicorn WSGI Server",
    "After=network.target",
    "",
    "[Service]",
    "User=ropaapp",
    "Group=ropaapp",
    "WorkingDirectory=/opt/ropa",
    "Environment=FLASK_ENV=production",
    "Environment=FLASK_SECRET_KEY=<gere-com-python-secrets.token_hex(32)>",
    "ExecStart=/opt/ropa/.venv/bin/gunicorn app:app \\",
    "    --workers 2 \\",
    "    --bind unix:/run/ropa/ropa.sock \\",
    "    --access-logfile /var/log/ropa/access.log \\",
    "    --error-logfile /var/log/ropa/error.log",
    "Restart=always",
    "RestartSec=5",
    "RuntimeDirectory=ropa",
    "",
    "[Install]",
    "WantedBy=multi-user.target",
])
add_para(doc, "Ativar e iniciar: sudo systemctl enable ropa && sudo systemctl start ropa", color=CINZA, italic=True)
add_space(doc)

add_heading(doc, "6.2 Configuração Nginx com TLS", level=2)
add_para(doc, "Criar /etc/nginx/sites-available/ropa e ativar com symlink:")
add_code_block(doc, [
    "server {",
    "    listen 443 ssl http2;",
    "    server_name ropa.intranet.example.org;",
    "",
    "    ssl_certificate     /etc/ssl/ropa/ropa.crt;",
    "    ssl_certificate_key /etc/ssl/ropa/ropa.key;",
    "    ssl_protocols       TLSv1.2 TLSv1.3;",
    "    ssl_ciphers         HIGH:!aNULL:!MD5;",
    "    ssl_session_cache   shared:SSL:10m;",
    "",
    "    # Headers de segurança HTTP",
    "    add_header Strict-Transport-Security 'max-age=31536000; includeSubDomains' always;",
    "    add_header X-Content-Type-Options    nosniff;",
    "    add_header X-Frame-Options           SAMEORIGIN;",
    "    add_header X-XSS-Protection          '1; mode=block';",
    "    add_header Referrer-Policy           'strict-origin-when-cross-origin';",
    "",
    "    # Rate limiting (100 req/min por IP)",
    "    limit_req_zone $binary_remote_addr zone=ropa:10m rate=100r/m;",
    "    limit_req zone=ropa burst=20 nodelay;",
    "",
    "    client_max_body_size 10M;",
    "",
    "    location / {",
    "        proxy_pass         http://unix:/run/ropa/ropa.sock;",
    "        proxy_set_header   Host              $host;",
    "        proxy_set_header   X-Real-IP         $remote_addr;",
    "        proxy_set_header   X-Forwarded-For   $proxy_add_x_forwarded_for;",
    "        proxy_set_header   X-Forwarded-Proto $scheme;",
    "        proxy_read_timeout 120s;",
    "    }",
    "}",
    "server {",
    "    listen 80;",
    "    server_name ropa.intranet.example.org;",
    "    return 301 https://$host$request_uri;",
    "}",
])
add_space(doc)

add_heading(doc, "6.3 Script de Backup Automático", level=2)
add_code_block(doc, [
    "#!/bin/bash",
    "# /opt/ropa/scripts/backup.sh",
    "# Cron: 0 2 * * * /opt/ropa/scripts/backup.sh",
    "",
    "BACKUP_DIR=/nas/backups/ropa",
    "DB=/opt/ropa/ropa.db",
    "LOG=/var/log/ropa/backup.log",
    "DATE=$(date +%Y%m%d_%H%M%S)",
    "",
    "mkdir -p $BACKUP_DIR",
    "",
    "# Backup quente do SQLite (seguro durante leitura/escrita)",
    'sqlite3 $DB ".backup $BACKUP_DIR/ropa_$DATE.db"',
    "",
    "# Retenção: manter apenas os últimos 30 backups",
    "ls -t $BACKUP_DIR/*.db | tail -n +31 | xargs -r rm",
    "",
    'echo "[$(date)] Backup: ropa_$DATE.db" >> $LOG',
])
add_space(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. CRONOGRAMA
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "7. Cronograma Resumido", level=1)

crono = [
    # (atividade, inicio, fim, resp, is_phase)
    ("FASE 0 — Preparação e Planejamento",          "Semana 1", "Semana 1",  "LGPD + Unidade de TI + SEGINF", True),
    ("  Kickoff e alinhamento entre unidades",       "Dia 1",    "Dia 2",     "Todas as unidades",     False),
    ("  Provisionamento VM + DNS interno",           "Dia 2",    "Dia 5",     "Unidade de TI",                 False),
    ("  Certificado TLS + regra de firewall",        "Dia 3",    "Dia 5",     "SEGINF",                False),
    ("FASE 1 — Ambiente de Homologação",             "Semana 2", "Semana 3",  "Unidade de TI + LGPD",         True),
    ("  Instalação e configuração do SO",            "Semana 2", "Semana 2",  "Unidade de TI",                 False),
    ("  Deploy e configuração em homologação",       "Semana 2", "Semana 3",  "Unidade de TI",                 False),
    ("  Testes funcionais e aceite da LGPD",        "Semana 3", "Semana 3",  "LGPD",                 False),
    ("FASE 2 — Segurança e Hardening",               "Semana 4", "Semana 4",  "SEGINF + Unidade de TI",        True),
    ("  Auditoria de segurança + RIPD",              "Semana 4", "Semana 4",  "SEGINF + LGPD",        False),
    ("FASE 3 — Implantação em Produção",             "Semana 5", "Semana 6",  "Unidade de TI + LGPD",         True),
    ("  Deploy em produção + treinamento",           "Semana 5", "Semana 6",  "Unidade de TI + LGPD",         False),
    ("  Termo de Aceite de Implantação",             "Semana 6", "Semana 6",  "LGPD + Unidade de TI + SEGINF",False),
    ("FASE 4 — Operação Contínua",                   "Semana 7+","Contínuo",  "Unidade de TI + LGPD",         True),
]
t = doc.add_table(rows=len(crono) + 1, cols=4)
make_header_cell(t.cell(0, 0), "Atividade / Marco")
make_header_cell(t.cell(0, 1), "Início", size=9)
make_header_cell(t.cell(0, 2), "Fim", size=9)
make_header_cell(t.cell(0, 3), "Responsável")
for i, (ativ, ini, fim, resp, is_phase) in enumerate(crono, 1):
    bg = AZUL_CLARO_HEX if is_phase else CINZA_CLARO_HEX
    make_data_cell(t.cell(i, 0), ativ, bg=bg, bold=is_phase, size=9)
    make_data_cell(t.cell(i, 1), ini,  bg=bg, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    make_data_cell(t.cell(i, 2), fim,  bg=bg, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    make_data_cell(t.cell(i, 3), resp, bg=bg, bold=is_phase, size=9)

add_space(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 8. MATRIZ RACI
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "8. Matriz de Responsabilidades (RACI)", level=1)
add_para(doc, "R = Responsável pela execução  ·  A = Aprovador  ·  C = Consultado  ·  I = Informado", color=CINZA)
add_space(doc, 4)

raci = [
    ("Provisionamento de infraestrutura",       "R", "I", "C", "A"),
    ("Instalação e configuração do sistema",    "R", "C", "C", "I"),
    ("Configuração de segurança / TLS",         "R", "I", "A", "I"),
    ("Testes funcionais da aplicação",          "C", "R", "C", "I"),
    ("Aprovação para ir a produção",            "C", "R", "A", "A"),
    ("Treinamento de usuários",                 "I", "R", "I", "C"),
    ("Backup e monitoramento contínuo",         "R", "I", "C", "I"),
    ("Manutenção de patches e atualizações",    "R", "I", "A", "I"),
    ("Revisão mensal do RoPA",                  "I", "R", "I", "I"),
    ("Elaboração e revisão do RIPD",            "I", "R", "A", "C"),
    ("Auditoria de acessos (semestral)",        "C", "R", "A", "I"),
]
t = doc.add_table(rows=len(raci) + 1, cols=5)
make_header_cell(t.cell(0, 0), "Atividade")
for j, unidade in enumerate(["Unidade de TI", "LGPD", "SEGINF", "SECAD"], 1):
    make_header_cell(t.cell(0, j), unidade, size=9)
for i, (ativ, *valores) in enumerate(raci, 1):
    bg = AZUL_CLARO_HEX if i % 2 == 1 else CINZA_CLARO_HEX
    make_data_cell(t.cell(i, 0), ativ, bg=bg, size=9)
    cores_raci = {"R": AZUL, "A": VERMELHO, "C": CINZA, "I": PRETO}
    for j, val in enumerate(valores, 1):
        make_data_cell(
            t.cell(i, j), val, bg=bg, size=9,
            bold=(val == "R"),
            text_color=cores_raci.get(val, PRETO),
            align=WD_ALIGN_PARAGRAPH.CENTER
        )

add_space(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 9. RISCOS E MITIGAÇÕES
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "9. Riscos e Mitigações", level=1)

riscos = [
    ("Atraso no provisionamento da VM",          "Média",  "Alto",     "Solicitar na Semana 0; manter servidor alternativo de contingência"),
    ("Falha no certificado TLS interno",         "Baixa",  "Alto",     "Solicitar certificado na Semana 0; testar em homologação antes da produção"),
    ("Conflito de versão Python no SO",          "Baixa",  "Médio",    "Usar pyenv ou compilar Python 3.10; isolar ambiente no venv"),
    ("Corrupção do banco SQLite",                "Baixa",  "Alto",     "Backup diário automatizado; restore test semestral; migrar para PostgreSQL a médio prazo"),
    ("Acesso não autorizado ao sistema",         "Média",  "Alto",     "Restringir por firewall à rede interna; implementar autenticação LDAP/SSO"),
    ("Perda de dados por falha de hardware",     "Baixa",  "Crítico",  "Backup em NAS separado; replicação para storage Instituição"),
    ("Resistência à adoção pelos usuários",      "Média",  "Médio",    "Treinamento; suporte da DPO; comunicação sobre obrigatoriedade LGPD Art. 37"),
    ("CVE em dependência Python",                "Média",  "Alto",     "pip-audit trimestral; monitorar CVE databases; política de atualização documentada"),
]
prob_cor  = {"Alta": VERMELHO, "Média": LARANJA, "Baixa": VERDE}
imp_cor   = {"Crítico": RGBColor(0x7F, 0x1D, 0x1D), "Alto": VERMELHO, "Médio": LARANJA, "Baixo": VERDE}

t = doc.add_table(rows=len(riscos) + 1, cols=4)
make_header_cell(t.cell(0, 0), "Risco")
make_header_cell(t.cell(0, 1), "Probabilidade")
make_header_cell(t.cell(0, 2), "Impacto")
make_header_cell(t.cell(0, 3), "Mitigação")
for i, (risco, prob, imp, mit) in enumerate(riscos, 1):
    bg = AZUL_CLARO_HEX if i % 2 == 1 else CINZA_CLARO_HEX
    make_data_cell(t.cell(i, 0), risco, bg=bg, size=9)
    make_data_cell(t.cell(i, 1), prob, bg=bg, size=9, bold=True, text_color=prob_cor.get(prob, PRETO), align=WD_ALIGN_PARAGRAPH.CENTER)
    make_data_cell(t.cell(i, 2), imp,  bg=bg, size=9, bold=True, text_color=imp_cor.get(imp, PRETO),   align=WD_ALIGN_PARAGRAPH.CENTER)
    make_data_cell(t.cell(i, 3), mit,  bg=bg, size=9)

add_space(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 10. CHECKLIST DE GO-LIVE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "10. Checklist de Go-Live", level=1)
add_para(doc, "Todos os itens abaixo devem estar concluídos antes da abertura do acesso em produção:")
add_space(doc, 4)

add_heading(doc, "Infraestrutura", level=2)
for item in [
    "VM de produção provisionada com Ubuntu Server 22.04 LTS e patches aplicados",
    "Usuário de sistema 'ropaapp' criado sem shell de login",
    "Permissões de diretório /opt/ropa verificadas (proprietário: ropaapp:ropaapp, 750)",
    "Serviço systemd ropa.service ativo e configurado para iniciar no boot (enabled)",
    "Nginx instalado, configurado e redirecionando HTTP → HTTPS corretamente",
    "Certificado TLS válido instalado e testado (verificar data de expiração)",
    "Firewall UFW configurado: portas 22 (SSH restrito por IP) e 443 abertas",
]:
    add_check(doc, item)

add_heading(doc, "Aplicação", level=2)
for item in [
    "Código-fonte clonado do repositório Git (branch main)",
    "Dependências instaladas no .venv: flask, gunicorn, openpyxl, reportlab",
    "Banco de dados inicializado sem erros (tabelas atividades e historico criadas)",
    "FLASK_SECRET_KEY definida como variável de ambiente (não estática no código)",
    "Exportação de XLSX e PDF testada com sucesso em produção",
    "Todos os endpoints retornando HTTP 200 (smoke test executado e documentado)",
]:
    add_check(doc, item)

add_heading(doc, "Segurança", level=2)
for item in [
    "Headers de segurança Nginx verificados (HSTS, X-Frame-Options, X-Content-Type-Options)",
    "Varredura com pip-audit executada — sem achados críticos abertos",
    "RIPD elaborado e aprovado pela LGPD (número do processo SEI registrado)",
    "Logs de acesso Nginx ativos e rotacionados pelo logrotate",
    "Backup automático configurado e restore testado com sucesso",
]:
    add_check(doc, item)

add_heading(doc, "Processos e Governança", level=2)
for item in [
    "Treinamento realizado com as unidades cadastradoras (lista de presença arquivada no SEI)",
    "URL de produção comunicada aos usuários autorizados",
    "Sistema registrado no inventário de ativos de TI do Instituição",
    "Runbook de operação e recuperação documentado e entregue à Unidade de TI",
    "Termo de Aceite de Implantação assinado por LGPD, Unidade de TI e SEGINF",
]:
    add_check(doc, item)

add_space(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 11. APROVAÇÃO
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "11. Aprovação e Assinaturas", level=1)
add_para(doc, (
    "Este documento é considerado aprovado após coleta das assinaturas das unidades responsáveis. "
    "O original deve ser arquivado como processo SEI no Instituição."
))
add_space(doc, 12)

aprovacao = [
    ("Unidade de Proteção de Dados — DPO",     "Encarregado(a) de Proteção de Dados"),
    ("Unidade de TI — Tecnologia da Informação",         ""),
    ("SEGINF — Segurança da Informação",         ""),
    ("SECAD — Assessoria de Administração",      ""),
]
t = doc.add_table(rows=len(aprovacao) + 1, cols=3)
make_header_cell(t.cell(0, 0), "Unidade")
make_header_cell(t.cell(0, 1), "Responsável / Nome")
make_header_cell(t.cell(0, 2), "Assinatura / Data")
for i, (unidade, nome) in enumerate(aprovacao, 1):
    bg = AZUL_CLARO_HEX if i % 2 == 1 else CINZA_CLARO_HEX
    make_data_cell(t.cell(i, 0), unidade, bg=bg, bold=True, size=9)
    make_data_cell(t.cell(i, 1), nome, bg=bg, size=9)
    # Célula de assinatura com espaço em branco generoso
    assin_cell = t.cell(i, 2)
    set_cell_bg(assin_cell, BRANCO_HEX)
    set_cell_borders(assin_cell)
    set_cell_margins(assin_cell, top=80, bottom=120, left=120, right=120)
    p_assin = assin_cell.paragraphs[0]
    r_data = p_assin.add_run("Data: _____ / _____ / 2026")
    r_data.font.size = Pt(9)
    r_data.font.name = "Arial"
    r_data.font.color.rgb = CINZA
    set_row_height(t.rows[i], 1.5)

add_space(doc, 36)

p_rod = doc.add_paragraph()
p_rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph_border_top(p_rod, color=AZUL_MEDIO_HEX, size=4)
p_rod.paragraph_format.space_before = Pt(12)
r_rod1 = p_rod.add_run(
    "Instituição  ·  Unidade de Proteção de Dados\n"
)
r_rod1.font.size = Pt(9)
r_rod1.font.name = "Arial"
r_rod1.font.color.rgb = CINZA

r_rod2 = p_rod.add_run(
    "Produzido nos termos da LGPD — Lei 13.709/2018, Art. 37  |  Versão 1.0 — Abril/2026"
)
r_rod2.italic = True
r_rod2.font.size = Pt(8)
r_rod2.font.name = "Arial"
r_rod2.font.color.rgb = CINZA


# ═══════════════════════════════════════════════════════════════════════════════
# SALVAR
# ═══════════════════════════════════════════════════════════════════════════════

output = Path(__file__).parent / "Roadmap_Implantacao_RoPA.docx"
doc.save(output)
print(f"OK → {output}")
